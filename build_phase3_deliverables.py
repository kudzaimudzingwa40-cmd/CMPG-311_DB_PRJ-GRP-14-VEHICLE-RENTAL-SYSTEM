from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT = Path("deliverables")
OUT.mkdir(exist_ok=True)

TITLE = "DriveFlow Rental"
SUBTITLE = "Phase 3 Physical Design Report"
DATE = "15 May 2026"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        shade_cell(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    for line in code.strip().splitlines():
        r = p.add_run(line.rstrip() + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(31, 58, 95)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Title", 24, "0B2545"),
        ("Subtitle", 13, "4B5563"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if "Heading" in name else 0)
        style.paragraph_format.space_after = Pt(6)
    return doc


sql_sections = [
    (
        "Core table creation example",
        """
CREATE TABLE IF NOT EXISTS vehicles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    vin TEXT UNIQUE NOT NULL,
    make TEXT NOT NULL,
    model TEXT NOT NULL,
    year INTEGER NOT NULL,
    license_plate TEXT UNIQUE NOT NULL,
    category TEXT NOT NULL CHECK (category IN ('Economy','SUV','Luxury')),
    daily_rate REAL NOT NULL CHECK (daily_rate > 0),
    status TEXT NOT NULL DEFAULT 'Available'
        CHECK (status IN ('Available','Rented','Maintenance')),
    mileage INTEGER DEFAULT 0 CHECK (mileage >= 0),
    last_service_mileage INTEGER DEFAULT 0 CHECK (last_service_mileage >= 0),
    next_service_date TEXT,
    image_url TEXT
);
""",
    ),
    (
        "Indexes",
        """
CREATE INDEX IF NOT EXISTS idx_vehicles_category_status
ON vehicles(category, status);

CREATE INDEX IF NOT EXISTS idx_bookings_vehicle_dates
ON bookings(vehicle_id, pickup_date, return_date);

CREATE INDEX IF NOT EXISTS idx_bookings_user_status
ON bookings(user_id, status);
""",
    ),
    (
        "Views",
        """
CREATE VIEW IF NOT EXISTS vw_booking_summary AS
SELECT b.id AS booking_id, u.name AS customer_name, u.email AS customer_email,
       v.make || ' ' || v.model AS vehicle_name, v.category,
       b.pickup_date, b.return_date, b.total_amount, b.discount_amount, b.status
FROM bookings b
JOIN users u ON u.id = b.user_id
JOIN vehicles v ON v.id = b.vehicle_id;
""",
    ),
    (
        "Query 1: vehicle search with LIKE, AND, OR, sorting, and row limitation",
        """
SELECT id, make, model, year, category, daily_rate
FROM vehicles
WHERE status != 'Maintenance'
  AND (make LIKE :search_term OR model LIKE :search_term)
ORDER BY category, daily_rate
LIMIT 20;
""",
    ),
    (
        "Query 2: date-based availability conflict check",
        """
SELECT id
FROM bookings
WHERE vehicle_id = :vehicle_id
  AND status NOT IN ('Cancelled','Returned')
  AND NOT (return_date <= :pickup_date OR pickup_date >= :return_date);
""",
    ),
    (
        "Query 3: revenue by category with aggregate, rounding, GROUP BY, and HAVING",
        """
SELECT v.category,
       COUNT(b.id) AS total_bookings,
       ROUND(SUM(b.total_amount), 2) AS total_revenue,
       ROUND(AVG(b.total_amount), 2) AS average_booking_value
FROM vehicles v
JOIN bookings b ON b.vehicle_id = v.id
WHERE b.status IN ('Confirmed','Returned')
GROUP BY v.category
HAVING COUNT(b.id) >= 1
ORDER BY total_revenue DESC;
""",
    ),
    (
        "Query 4: customer penalty balance with join and sub-query",
        """
SELECT u.name, u.email,
       (SELECT COALESCE(SUM(p.amount), 0)
        FROM penalties p
        JOIN bookings b2 ON b2.id = p.booking_id
        WHERE b2.user_id = u.id AND p.status = 'Unpaid') AS outstanding_balance
FROM users u
WHERE u.role = 'customer'
ORDER BY outstanding_balance DESC
LIMIT 10;
""",
    ),
    (
        "Query 5: date functions for maintenance due soon",
        """
SELECT make, model, license_plate, next_service_date
FROM vehicles
WHERE next_service_date IS NOT NULL
  AND next_service_date <= date('now', '+30 days')
ORDER BY next_service_date ASC;
""",
    ),
    (
        "Query 6: character functions and variable-style parameter for promo validation",
        """
SELECT code, description, discount_type, discount_value
FROM promo_codes
WHERE UPPER(code) = UPPER(TRIM(:promo_code))
  AND active = 1
  AND uses < max_uses
  AND (expires_at IS NULL OR expires_at > datetime('now'));
""",
    ),
]


def build_sql_file():
    db_text = Path("db.py").read_text(encoding="utf-8")
    schema_match = re.search(r'c\.executescript\("""(.*?)"""\)', db_text, re.S)
    schema_sql = schema_match.group(1).strip() if schema_match else "-- Schema block could not be extracted from db.py."
    parts = [
        "-- DriveFlow Rental Phase 3 SQL Evidence",
        "-- Raw SQL statements for physical design, database objects, and query requirements.",
        "",
        "-- Complete schema block from db.py: tables, constraints, indexes, and views",
        "",
        schema_sql,
        "",
    ]
    parts.append("-- Query evidence mapped to Phase 3 requirements")
    parts.append("")
    for title, sql in sql_sections:
        parts.append(f"-- {title}")
        parts.append(sql.strip())
        parts.append("")
    Path(OUT / "DriveFlow_Rental_Phase3_SQL_Evidence.sql").write_text("\n".join(parts), encoding="utf-8")


def build_demo_checklist():
    text = """# DriveFlow Rental Phase 3 Demonstration Checklist

Use this checklist when recording the Phase 3 video demonstration.

1. Open the project and start the Flask application with `python logic.py`.
2. Show that the SQLite database is initialised through raw SQL in `db.py`.
3. Demonstrate tables, primary keys, foreign keys, check constraints, indexes, and views.
4. Register or log in as a customer.
5. Search for vehicles using category and text filtering.
6. Open a vehicle detail page and create a booking with valid dates.
7. Apply a promo code or loyalty option where applicable.
8. Complete payment and show the generated receipt.
9. Log in as admin and show dashboard, fleet, bookings, customers, reports, promos, maintenance, penalties, and audit log.
10. Demonstrate at least one report/export and explain the query behind it.
11. Show extra functionality: loyalty points, promotions, penalties, maintenance logs, receipts, admin reports, and branded interface.
12. Close with a summary of how the system satisfies the Phase 3 physical design rubric.
"""
    Path(OUT / "DriveFlow_Rental_Phase3_Demo_Checklist.md").write_text(text, encoding="utf-8")


def build_docx():
    doc = setup_doc()

    title = doc.add_paragraph(TITLE, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(SUBTITLE, style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(f"Group 14 | {DATE}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(
        "This report documents the physical database design and SQL implementation evidence for the DriveFlow Rental vehicle rental system."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "DriveFlow Rental is a digital vehicle rental platform that supports customer registration, vehicle browsing, booking, payment, receipts, loyalty rewards, penalties, maintenance tracking, reporting, and administrative fleet control. Phase 3 focuses on the physical implementation of the database: tables, keys, constraints, indexes, views, data population, and SQL queries that support the company information requirements."
    )
    doc.add_paragraph(
        "The implementation uses raw SQL through SQLite. Database operations are written with direct SQL statements in the application data layer and route logic, without an object-relational mapper."
    )

    doc.add_heading("Rubric Alignment", level=1)
    add_table(
        doc,
        ["Rubric area", "Evidence in this deliverable"],
        [
            ["Database Objects", "Tables, data types, primary keys, foreign keys, check constraints, indexes, views, and seed data are documented."],
            ["Queries", "A query matrix and SQL evidence file demonstrate filtering, limitations, sorting, LIKE/AND/OR, functions, aggregates, grouping, joins, and sub-queries."],
            ["Video Demonstration", "A recording checklist is included to guide the demonstration of database setup, SQL statements, application workflows, and extra functionality."],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("1. Physical Database Objects", level=1)
    doc.add_paragraph(
        "The physical design converts the logical model into implemented SQL objects. Each table uses a numeric primary key, foreign keys for relationships, and constraints that protect controlled values and valid monetary/date ranges."
    )
    add_table(
        doc,
        ["Table", "Purpose", "Key constraints"],
        [
            ["users", "Stores customers, staff, and administrators.", "Unique email; role check; loyalty points cannot be negative."],
            ["vehicles", "Stores rentable vehicles and fleet state.", "Unique VIN and plate; category/status checks; positive daily rate."],
            ["bookings", "Stores rental reservations.", "FK to user and vehicle; valid date range; status check; non-negative totals."],
            ["payments", "Stores completed or pending payment records.", "FK to booking; unique payment reference; non-negative amount."],
            ["penalties", "Stores late/damage/other penalties.", "FK to booking; status check; non-negative amount and days late."],
            ["returns", "Stores return processing records.", "FK to booking; condition check; non-negative return mileage."],
            ["reviews", "Stores customer ratings and comments.", "FKs to booking, user, vehicle; rating check between 1 and 5."],
            ["promo_codes", "Stores discount codes.", "Unique code; discount type check; active flag check; usage limits."],
            ["maintenance_logs", "Stores service and repair records.", "FK to vehicle and logged-by user; non-negative costs and mileage."],
            ["audit_logs", "Stores administrative and system activity.", "Optional FK to user; action/entity trace fields."],
            ["loyalty_transactions", "Stores loyalty earning, redemption, refund, and bonus records.", "FK to user; optional FK to booking."],
            ["waitlist", "Stores customer waitlist requests.", "FKs to user and vehicle; waitlist status check."],
        ],
        [1.35, 2.25, 2.9],
    )

    doc.add_heading("Indexes", level=2)
    add_table(
        doc,
        ["Index", "Reason"],
        [
            ["idx_users_role", "Speeds up customer/admin/staff filtering."],
            ["idx_vehicles_category_status", "Supports vehicle search by category and availability."],
            ["idx_bookings_user_status", "Supports customer dashboards and booking status filtering."],
            ["idx_bookings_vehicle_dates", "Supports date-overlap availability checks."],
            ["idx_payments_booking", "Speeds up receipt/payment lookup by booking."],
            ["idx_penalties_booking_status", "Supports outstanding penalty lookups."],
            ["idx_reviews_vehicle", "Supports vehicle rating summaries."],
            ["idx_maintenance_vehicle_date", "Supports service history and due-service reports."],
        ],
        [2.4, 4.1],
    )

    doc.add_heading("Views", level=2)
    add_table(
        doc,
        ["View", "Purpose"],
        [
            ["vw_available_vehicles", "Reusable view of vehicles currently available to customers."],
            ["vw_booking_summary", "Combines booking, customer, and vehicle details for reporting."],
            ["vw_revenue_by_category", "Aggregates booking revenue by vehicle category."],
            ["vw_customer_penalty_balance", "Summarises outstanding customer penalties."],
        ],
        [2.4, 4.1],
    )

    doc.add_heading("Data Population", level=2)
    doc.add_paragraph(
        "The database initialisation script seeds administrator and staff users, a fleet of vehicles across Economy, SUV, and Luxury categories, and default promotional codes. This gives the application enough test data for booking, payment, reporting, and maintenance workflows."
    )

    doc.add_heading("2. SQL Query Evidence", level=1)
    doc.add_paragraph(
        "The following matrix shows how the SQL evidence satisfies the query requirements. The full SQL statements are included in the separate SQL evidence file."
    )
    add_table(
        doc,
        ["Requirement", "Implemented example"],
        [
            ["Company information requirements", "Vehicle search, booking availability, revenue reports, penalty balances, maintenance due-soon report."],
            ["Row and column limitations", "LIMIT clauses and selected column lists are used for dashboards and reports."],
            ["Sorting operations", "ORDER BY is used for price, category, date, and revenue orderings."],
            ["LIKE, AND, OR", "Vehicle search uses LIKE with AND/OR conditions."],
            ["Variables and character functions", "Named parameters with UPPER and TRIM are used for promo validation."],
            ["Rounding/truncation", "ROUND is used for revenue and average booking value calculations."],
            ["Date functions", "date, datetime, strftime, and julianday support validation and reports."],
            ["Aggregate functions", "COUNT, SUM, AVG, and COALESCE are used in reporting."],
            ["GROUP BY and HAVING", "Revenue by category groups bookings and filters grouped results."],
            ["Joins", "Booking summaries join bookings, users, vehicles, payments, and penalties."],
            ["Sub-queries", "Customer penalty balances and review counts use sub-queries."],
        ],
        [2.4, 4.1],
    )

    for title_text, sql in sql_sections[3:]:
        doc.add_heading(title_text, level=2)
        add_code(doc, sql)

    doc.add_heading("3. Extra Functionalities", level=1)
    add_bullets(
        doc,
        [
            "Customer loyalty points for welcome bonuses, payments, reviews, and clean returns.",
            "Promo code validation with percentage and fixed-amount discount support.",
            "Deposit support for bookings.",
            "Customer dashboards with booking history, unpaid penalties, and loyalty history.",
            "Payment receipts and administrative invoice generation.",
            "Admin reporting, CSV exports, maintenance records, audit logs, and customer management.",
            "Staff workflow for vehicle status updates and maintenance logging.",
            "Branded DriveFlow Rental interface with vehicle imagery and a custom logo.",
        ],
    )

    doc.add_heading("4. Demonstration Video Plan", level=1)
    add_numbered(
        doc,
        [
            "Start the application and show that the database initialises successfully.",
            "Open the SQL schema and explain tables, keys, constraints, indexes, and views.",
            "Demonstrate vehicle search and category filtering.",
            "Create a booking, complete payment, and open the receipt.",
            "Show the customer dashboard, penalties, profile, and review flow.",
            "Log in as admin and demonstrate fleet, bookings, customers, promos, reports, maintenance, penalties, and audit logs.",
            "Show at least one CSV export and explain the SQL query behind it.",
            "Conclude by mapping the implemented features back to the Phase 3 requirements.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The Phase 3 implementation provides a complete physical database design for DriveFlow Rental. The system includes normalized raw SQL tables, relationship enforcement, constraints, indexes, views, populated sample data, and SQL queries that support customer, staff, and administrator workflows. The application demonstrates both database correctness and practical functionality through booking, payment, reporting, loyalty, promotion, maintenance, and audit features."
    )

    doc.add_heading("Appendix A: Deliverable Files", level=1)
    add_bullets(
        doc,
        [
            "DriveFlow_Rental_Phase3_Physical_Design_Report.docx",
            "DriveFlow_Rental_Phase3_Physical_Design_Report.pdf",
            "DriveFlow_Rental_Phase3_SQL_Evidence.sql",
            "DriveFlow_Rental_Phase3_Demo_Checklist.md",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "DriveFlow Rental - Phase 3 Physical Design"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT / "DriveFlow_Rental_Phase3_Physical_Design_Report.docx")


def pdf_table(headers, rows, widths):
    data = [headers] + rows
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#C8D3DF")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf():
    pdf_path = OUT / "DriveFlow_Rental_Phase3_Physical_Design_Report.pdf"
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="DriveFlow Rental Phase 3 Physical Design Report",
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontSize=25,
            leading=30,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSub",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Blue",
            parent=styles["Heading1"],
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2Blue",
            parent=styles["Heading2"],
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyTight",
            parent=styles["BodyText"],
            fontSize=9.5,
            leading=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeBlock",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=7.4,
            leading=9,
            textColor=colors.HexColor("#1F3A5F"),
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#D8E0EA"),
            borderWidth=0.4,
            borderPadding=5,
            spaceAfter=7,
        )
    )

    story = [
        Spacer(1, 1.4 * inch),
        Paragraph(TITLE, styles["CoverTitle"]),
        Paragraph(SUBTITLE, styles["CoverSub"]),
        Paragraph(f"Group 14 | {DATE}", styles["CoverSub"]),
        Spacer(1, 0.55 * inch),
        Paragraph(
            "Physical database design, SQL implementation evidence, query coverage, extra functionality, and demonstration guidance.",
            styles["CoverSub"],
        ),
        PageBreak(),
        Paragraph("Executive Summary", styles["H1Blue"]),
        Paragraph(
            "DriveFlow Rental is a digital vehicle rental platform supporting customer registration, vehicle browsing, booking, payment, receipts, loyalty rewards, penalties, maintenance tracking, reporting, and administrative fleet control. This Phase 3 report documents the physical database implementation using raw SQL through SQLite.",
            styles["BodyTight"],
        ),
        Paragraph("Rubric Alignment", styles["H1Blue"]),
        pdf_table(
            ["Rubric area", "Evidence"],
            [
                ["Database Objects", "Tables, data types, primary keys, foreign keys, check constraints, indexes, views, and populated data."],
                ["Queries", "SQL evidence for filtering, limits, sorting, LIKE/AND/OR, functions, aggregates, grouping, joins, and sub-queries."],
                ["Video Demonstration", "Checklist for showing setup, working SQL statements, application workflows, and extra functionality."],
            ],
            [1.7 * inch, 4.8 * inch],
        ),
        Spacer(1, 8),
        Paragraph("1. Physical Database Objects", styles["H1Blue"]),
        Paragraph(
            "The design implements twelve main database tables. Controlled fields use CHECK constraints, relationships use foreign keys, and lookup-heavy workflows are supported with indexes and views.",
            styles["BodyTight"],
        ),
        pdf_table(
            ["Table", "Purpose", "Key constraints"],
            [
                ["users", "Customers, staff, and administrators.", "Unique email; role check; non-negative loyalty points."],
                ["vehicles", "Rentable vehicle fleet.", "Unique VIN/plate; category/status checks; positive daily rate."],
                ["bookings", "Rental reservations.", "FKs to users/vehicles; valid date range; status and amount checks."],
                ["payments", "Payment records.", "FK to booking; unique reference; non-negative amount."],
                ["penalties", "Penalty records.", "FK to booking; status, amount, and days-late checks."],
                ["returns", "Return processing.", "FK to booking; condition and mileage checks."],
                ["reviews", "Customer ratings.", "Rating must be between 1 and 5."],
                ["promo_codes", "Discount codes.", "Unique code; type, active, and usage checks."],
                ["maintenance_logs", "Service history.", "FK to vehicle; non-negative cost and mileage."],
                ["audit_logs", "Activity tracking.", "Action/entity trace with optional user FK."],
                ["loyalty_transactions", "Points history.", "FK to user and optional booking FK."],
                ["waitlist", "Waitlist requests.", "FKs to user/vehicle; waitlist status check."],
            ],
            [1.25 * inch, 2.15 * inch, 3.1 * inch],
        ),
        Paragraph("Indexes and Views", styles["H2Blue"]),
        pdf_table(
            ["Object", "Purpose"],
            [
                ["idx_vehicles_category_status", "Fast vehicle category and availability filtering."],
                ["idx_bookings_vehicle_dates", "Fast availability conflict checks."],
                ["vw_available_vehicles", "Reusable list of customer-bookable vehicles."],
                ["vw_booking_summary", "Booking, customer, and vehicle report view."],
                ["vw_revenue_by_category", "Aggregated category revenue report."],
                ["vw_customer_penalty_balance", "Outstanding customer penalty summary."],
            ],
            [2.25 * inch, 4.25 * inch],
        ),
        Paragraph("2. SQL Query Evidence", styles["H1Blue"]),
        Paragraph(
            "The full SQL evidence script accompanies this report. Representative statements are shown below.",
            styles["BodyTight"],
        ),
    ]

    for title_text, sql in sql_sections[3:]:
        story.append(Paragraph(title_text, styles["H2Blue"]))
        story.append(Preformatted(sql.strip(), styles["CodeBlock"]))

    story.extend(
        [
            Paragraph("3. Extra Functionalities", styles["H1Blue"]),
            Paragraph(
                "Extra functionality includes loyalty points, promo codes, deposits, receipts, invoice generation, customer dashboards, admin reporting, CSV exports, penalties, maintenance logs, audit logs, staff fleet updates, branded UI, and a custom DriveFlow Rental logo.",
                styles["BodyTight"],
            ),
            Paragraph("4. Demonstration Video Plan", styles["H1Blue"]),
        ]
    )
    for item in [
        "Start the application with python logic.py and show successful database initialisation.",
        "Explain tables, keys, constraints, indexes, and views in the SQL/database layer.",
        "Demonstrate customer registration, vehicle search, booking, payment, and receipt.",
        "Show the customer dashboard, penalties, profile, and review workflow.",
        "Log in as admin and show fleet, bookings, customers, promos, reports, maintenance, penalties, and audit logs.",
        "Show at least one CSV export and explain the SQL query behind it.",
        "Conclude by mapping the implemented functionality back to the Phase 3 requirements.",
    ]:
        story.append(Paragraph(f"- {item}", styles["BodyTight"]))

    story.extend(
        [
            Paragraph("Conclusion", styles["H1Blue"]),
            Paragraph(
                "The Phase 3 implementation provides a complete physical database design for DriveFlow Rental. It includes normalized raw SQL tables, constraints, indexes, views, populated sample data, and SQL queries that support customer, staff, and administrator workflows.",
                styles["BodyTight"],
            ),
            Paragraph("Appendix A: Files Included", styles["H1Blue"]),
            Paragraph("DriveFlow_Rental_Phase3_Physical_Design_Report.docx", styles["BodyTight"]),
            Paragraph("DriveFlow_Rental_Phase3_Physical_Design_Report.pdf", styles["BodyTight"]),
            Paragraph("DriveFlow_Rental_Phase3_SQL_Evidence.sql", styles["BodyTight"]),
            Paragraph("DriveFlow_Rental_Phase3_Demo_Checklist.md", styles["BodyTight"]),
        ]
    )

    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#4B5563"))
        canvas.drawCentredString(letter[0] / 2, 0.42 * inch, f"DriveFlow Rental - Phase 3 Physical Design | Page {_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    build_sql_file()
    build_demo_checklist()
    build_docx()
    build_pdf()
    print("Phase 3 deliverables generated in", OUT.resolve())
